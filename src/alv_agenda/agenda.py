from docx import Document
from datetime import date
from .agenda_punt import AgendaPunt


class Agenda:
    """Agenda class."""
    title = "Algemene Leden Vergadering van het \"S. G. William Froude\""

    def __init__(self, agenda_punten, location, alv=True):
        """Hierin wordt de titel en de locatie en de tijd neergezet in het word
        bestand.
        Voor nu bestaat dit alleen uit een header.
        """
        self.agenda_punten = agenda_punten
        self.date = date.today().strftime("%d-%m-%Y")
        self.doc = Document()
        self.location = location
        self.standard_alv = alv

        # Create header
        with open('snippets/header.txt') as h:
            header = h.read()
            compiled_header = compile(header, '<header>', 'eval')
            evaluated_header = eval(compiled_header)
        self.doc.add_paragraph(evaluated_header)

        # Create standard points
        if self.standard_alv:
            alv()

        def alv(self):
            date_last_alv = input("Wanneer was de vorige ALV?")
            standard_punten = [
                    "Opening",
                    "Mededelingen",
                    "Mededelingen Onderwijs",
                    "Mededelingen Financieel",
                    f"Goedkeuring Notulen van ALV gehouden op {date_last_alv}"]
            for title in standard_punten:
                self.doc.add_paragraph(title, style='List Number')

    def add_agenda_punt(self, title, id, tijdsduur, begintijd):
        """Maak een agendapunt."""
        agenda_punt = AgendaPunt(title, id, tijdsduur, begintijd)
        self.agenda_punten.insert(agenda_punt.id, agenda_punt)

    def get_agenda_input(self):
        switch = True
        n = 1
        while switch:
            print(f"Adding point number {n}")
            title = input("What is the title? ")
            id = input("What is the id? ")
            tijdsduur = input("What is the tijdsduur? ")
            begintijd = input("What is the begintijd? ")
            self.add_agenda_punt(title, id, tijdsduur, begintijd)
            switch = input("Do you want to add another? [y/n] ").lower() == 'y'
            n += 1
        return "Done"

    def create_agenda(self):
        for agenda_punt in self.agenda_punten:
            self.doc.add_paragraph(agenda_punt.title, style='List Number')
        return "Agenda updated"

    def save_agenda(self):
        """Save the agenda."""
        self.doc.save('Agenda.docx')
        return True

    @classmethod
    def from_prompt(cls):
        """Build the agenda and save it."""
        # Todo: add input for agenda_punten.
        location = input("Where is the meeting being held? ")
        alv_agenda = cls([], location)
        alv_agenda.get_agenda_input()
        alv_agenda.create_agenda()
        alv_agenda.save_agenda()
        return True


class LangeAgenda(Agenda):
    """Uitbreiding op de normale Agenda"""

    def __init__(self, agenda_punten, location):
        super.__init__(agenda_punten, location)

    def uitgewerkt_stuk(self):
        """Maak een stuk voor de lange agenda."""
        # opties = [mededeling, begroting, exploitatie, presentatie]
        # if match met optie, dan pak optie-snippet, met invulling van
        # hierboven
        return None
