"""Create Agenda's."""
from docx import Document
from datetime import date
from agenda_punt import AgendaPunt
from import_agenda import AgendaData
import os

cwd = os.getcwd()


class AgendaBuilder:
    """Agenda builder class."""
    title = "Algemene Leden Vergadering van het S. G. \"William Froude\""
    document = Document()

    def __init__(self, agenda_points, location):
        """Initialize agenda. """
        self.agenda_points = agenda_points
        self.date = date.today().strftime("%d/%m/%Y")
        self.location = location

        # Add first things to Document
        self.document

    @classmethod
    def from_agenda_data(cls, workbook_location, location):
        """Create agenda points."""
        data_file = AgendaData(workbook_location)
        data = data_file.import_from_excel()
        agenda_point_list = []
        for id, data_list in data:
            agenda_point = AgendaPunt(id, data_list['type'],
                                      data_list['title'],
                                      data_list['duration'],
                                      data_list['speaker'],
                                      data_list['subject'])
            agenda_point_list.append(agenda_point)
        return cls(agenda_point_list, location)

    def write_paragraph(self, text: str):
        self.document.add_paragraph(text)
        return True

    def write_all_points(self):
        for point in self.agenda_points:
            self.write_paragraph(point.numbered_title())

    # def header(self):
    #     """Create header."""
    #     evaluated_header = eval(compile_paragraph(
    #         'alv_agenda/snippets/header.txt'))
    #     print(evaluated_header)
    #     self.doc.add_paragraph(evaluated_header)

    # def alv(self):
    #     """Bouw standaard punten voor ALV."""
    #     date_last_alv = input("Wanneer was de vorige ALV?")
    #     standard_punten = [
    #             "Opening",
    #             "Mededelingen",
    #             "Mededelingen Onderwijs",
    #             "Mededelingen Financieel",
    #             f"Goedkeuring Notulen van ALV gehouden op {date_last_alv}"]
    #     for title in standard_punten:
    #         self.doc.add_paragraph(title, style='List Number')

    # def add_agenda_punt(self, title, id, tijdsduur, begintijd):
    #     """Maak een agendapunt."""
    #     agenda_punt = AgendaPunt(title, id, tijdsduur, begintijd)
    #     self.agenda_punten.insert(agenda_punt.id, agenda_punt)

    # # def get_agenda_input(self):
    # #     switch = True
    # #     n = 1
    # #     while switch:
    # #         print(f"Adding point number {n}")
    # #         title = input("What is the title? ")
    # #         id = input("What is the id? ")
    # #         tijdsduur = input("What is the tijdsduur? ")
    # #         begintijd = input("What is the begintijd? ")
    # #         self.add_agenda_punt(title, id, tijdsduur, begintijd)
    # #         switch = input("Do you want to add another? [y/n] ").lower() == 'y'
    # #         n += 1
    # #     return "Done"

    # def create_agenda(self):
    #     for agenda_punt in self.agenda_punten:
    #         self.doc.add_paragraph(agenda_punt.title, style='List Number')
    #     return "Agenda updated"

    # def save_agenda(self):
    #     """Save the agenda."""
    #     self.doc.save('Agenda.docx')
    #     return True

    # @classmethod
    # def from_prompt(cls):
    #     """Build the agenda and save it."""
    #     # Todo: add input for agenda_punten.
    #     location = input("Where is the meeting being held? ")
    #     alv_agenda = cls([], location)
    #     alv_agenda.get_agenda_input()
    #     alv_agenda.create_agenda()
    #     alv_agenda.save_agenda()
    #     return True


# class LangeAgenda(Agenda):
    # """Uitbreiding op de normale Agenda."""

    # def __init__(self, agenda_punten, location):
    #     super().__init__(self, agenda_punten, location)

    # def header(self):
    #     """Create header."""
    #     evaluated_header = evaluate_paragraph('snippets/header_lang.txt')
    #     self.doc.add_paragraph(evaluated_header)

    # def alv(self):
    #     """Bouw standaard punten voor ALV."""
    #     date_last_alv = input("Wanneer was de vorige ALV?")
    #     standard_punten = [
    #             # Opening
    #             AgendaPunt("Opening", 1, 5, '19:03', text_location="opening"),

    #             # Mededelingen
    #             AgendaPunt("Mededelingen", 2, 5, '19:03',
    #                        text_location="mededelingen"),
    #             # Onderwijs
    #             AgendaPunt("Mededelingen Onderwijs", 3, 5,
    #                        '19:03', text_location="onderwijs"),

    #             # Financieel
    #             AgendaPunt("Mededelingen Financieel", 4, 5,
    #                        '19:03', text_location="financieel"),
    #             # Goedkeuring Notulen
    #             AgendaPunt(
    #                 f"Goedkeuring Notulen van ALV gehouden op {date_last_alv}",
    #                 5, 5, '19:03', text_location="notulen")
    #     ]
    #     for title in standard_punten:
    #         self.doc.add_heading(title, level=1, style='List Number')
    #         par = evaluate_paragraph(f'snippets/lange_agenda/{title}')
    #         self.doc.add_paragraph(par)


# def compile_paragraph(path):
    # with open(path) as p:
    #     paragraph = p.read()
    #     compiled_paragraph = compile(paragraph, f'<{paragraph}>', 'eval')
    # return compiled_paragraph

    # # def match_stuk(self, agendapunt):
    # #     """Maak een stuk voor de lange agenda."""
    # #     match agendapunt.type:
    # #         case "mededeling":
    # #         case "":
    # #         case "exploitatie":
    # #         case "presentatie":

    #     # opties = [mededeling, begroting, exploitatie, presentatie]
    #     # if match met optie, dan pak optie-snippet, met invulling van
    #     # hierboven
    #     # return None


# def main():
    # _1 = AgendaPunt("Opening", 1, 10, "19:03")


# if __name__ == "__main__":
    # main()
