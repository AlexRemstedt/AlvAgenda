"""Main document."""


from docx import Document
from datetime import date, time


class AgendaPunt:
    """AgendaPunt"""

    def __init__(self, title, nummer, tijdsduur, begintijd):
        self.title = title
        self.tijdsduur = tijdsduur
        self.nummer = nummer - 1
        [h, m] = [int(n) for n in begintijd.split(":")]
        self.begintijd = time(h, m)
        self.eindtijd = self.begintijd.minute + tijdsduur


class Agenda:
    """Agenda class."""
    title = "Algemene Leden Vergadering van het \"S. G. William Froude\""

    def __init__(self, agenda_punten, location):
        self.agenda_punten = agenda_punten
        self.date = date.today().strftime("%d-%m-%Y")
        self.doc = Document()
        self.location = location

        # Start met de agenda bouwen.
        self.init_agenda()

    def init_agenda(self):
        """Initiate agenda.

        Hierin wordt de titel en de locatie en de tijd neergezet in het word
        bestand.
        Voor nu bestaat dit alleen uit een header.
        """
        # Create header
        with open('snippets/header.txt') as h:
            header = h.read()
            compiled_header = compile(header, '<header>', 'eval')
            evaluated_header = eval(compiled_header)
        self.doc.add_paragraph(evaluated_header)

    def add_agenda_punt(self, title, nummer, tijdsduur, begintijd):
        """Maak een agendapunt."""
        agenda_punt = AgendaPunt(title, nummer, tijdsduur, begintijd)
        self.agenda_punten.insert(agenda_punt.nummer, agenda_punt)

    def create_agenda(self):
        for agenda_punt in self.agenda_punten:
            self.doc.add_paragraph(agenda_punt.title, style='List Number')
        return "Agenda updated"

    def save_agenda(self):
        """Save the agenda."""
        self.doc.save('Agenda.docx')
        return True

    def build_agenda(self):
        """Build the agenda and save it."""
        # Todo: add input for agenda_punten.
        self.create_agenda()
        self.save_agenda()


class LangeAgenda(Agenda):
    """Uitbreiding op de normale Agenda"""

    def __init__(self, agenda_punten, location):
        super.__init__(agenda_punten, location)

    def uitgewerkt_stuk(self):
        """Maak een stuk voor de lange agenda."""
        # opties = [mededeling, begroting, exploitatie, presentatie]
        # if match met optie, dan pak optie-snippet, met invulling van
        # hierboven
        return None


def main():
    alv_agenda = Agenda([], 'Hier')
    alv_agenda.add_agenda_punt("MMD", 2, 10, "19:03")
    alv_agenda.add_agenda_punt("DWT", 1, 10, "19:03")
    alv_agenda.create_agenda()
    alv_agenda.save_agenda()


if __name__ == "__main__":
    main()
